import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_heading(doc, text, level, align=WD_ALIGN_PARAGRAPH.LEFT):
    h = doc.add_heading(text, level=level)
    h.alignment = align
    return h

def add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, justify=False):
    p = doc.add_paragraph(text)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        p.alignment = align
    return p

def create_report():
    doc = Document()
    
    # Title Page
    doc.add_heading('INTERTALKS', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('Policy Compliance RAG System', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n\n\n')
    p = doc.add_paragraph('A Project Report\nSubmitted in partial fulfillment of the requirements\nfor the degree of\nBachelor of Technology')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n\n')
    doc.add_paragraph('Submitted by:\nNikhil Chauhan').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Abstract
    add_heading(doc, 'Abstract', 1, WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, 
        "This project presents 'InterTalks', a production-grade Retrieval-Augmented Generation (RAG) system specifically designed for policy compliance and document querying. "
        "Unlike generic conversational AI models, InterTalks provides strictly evidence-backed responses, eliminating hallucinations by operating within a bounded context. "
        "The system incorporates advanced document processing techniques, including section-based chunking rather than arbitrary token limits, ensuring semantic boundaries are respected. "
        "It integrates a multi-stage retrieval pipeline utilizing dense embeddings via Sentence Transformers and a re-ranking phase using a Cross-Encoder to achieve high precision. "
        "Governance is central to the architecture, featuring role-based access control (Admin, Reviewer, Contributor, Viewer), version control, and approval workflows. "
        "Only approved policies are indexed in the vector database (Qdrant), guaranteeing that all generated answers are based on the latest, authorized corporate information. "
        "When insufficient evidence is found, the system deterministically refuses to answer, prioritizing accuracy over fluency. "
        "The backend is powered by FastAPI and PostgreSQL, while the frontend offers a responsive, dark-themed React UI. "
        "This report details the system architecture, mathematical foundations of the embedding strategy, implementation specifics, and evaluation of the RAG pipeline.", justify=True)
    doc.add_page_break()

    # Introduction
    add_heading(doc, 'Chapter 1: Introduction', 1)
    add_heading(doc, '1.1 Problem Statement', 2)
    add_paragraph(doc, 
        "In modern enterprises, navigating extensive compliance documents, HR policies, and legal frameworks is a time-consuming challenge. "
        "Employees often struggle to find specific clauses or rules within hundreds of pages of documentation. "
        "Traditional keyword search mechanisms often fail to capture semantic intent, leading to irrelevant results or missed information. "
        "While Large Language Models (LLMs) offer natural language querying, using them out-of-the-box introduces the severe risk of 'hallucinations'—the generation of plausible but factually incorrect information. "
        "In compliance and legal contexts, incorrect answers can lead to severe operational and legal consequences. "
        "Therefore, there is a critical need for an intelligent system that can understand natural language queries but strictly restrict its answers to a curated, approved set of documents.", justify=True)
    
    add_heading(doc, '1.2 Objectives', 2)
    add_paragraph(doc, "The primary objectives of the InterTalks system are:", justify=True)
    doc.add_paragraph("1. To develop a Retrieval-Augmented Generation (RAG) pipeline that guarantees evidence-backed responses.", style='List Number')
    doc.add_paragraph("2. To implement a deterministic refusal mechanism when the system cannot find sufficient approved evidence.", style='List Number')
    doc.add_paragraph("3. To construct a rigorous document ingestion pipeline featuring section-based chunking rather than arbitrary token counts, maintaining logical coherence.", style='List Number')
    doc.add_paragraph("4. To enforce strict governance through an approval workflow, ensuring only 'Approved' document versions are queryable.", style='List Number')
    doc.add_paragraph("5. To design an intuitive, enterprise-grade user interface for querying, document management, and administration.", style='List Number')

    add_heading(doc, '1.3 Scope', 2)
    add_paragraph(doc, 
        "The scope of this project encompasses the end-to-end development of a web application and its underlying AI infrastructure. "
        "It includes parsing complex PDF and DOCX files to extract text and structure, generating dense vector representations of textual chunks, "
        "and utilizing Qdrant for fast similarity search. The application includes a React-based frontend tailored for different user roles. "
        "The system handles short query expansion, adaptive similarity thresholding, and composite scoring for re-ranking. "
        "It focuses on internal corporate policy querying and does not extend to general web search or unverified external data sources.", justify=True)
    doc.add_page_break()

    # Literature Review
    add_heading(doc, 'Chapter 2: Literature Review', 1)
    add_paragraph(doc, 
        "The evolution of natural language processing has seen a paradigm shift from traditional Information Retrieval (IR) methods (like BM25) to dense semantic search powered by transformers. "
        "Retrieval-Augmented Generation (RAG), introduced by Lewis et al. (2020), bridges the gap between parametric memory (internal LLM knowledge) and non-parametric memory (external databases). "
        "By retrieving relevant context and appending it to the LLM's prompt, RAG reduces hallucinations and allows the model to leverage up-to-date information without retraining. "
        "However, standard RAG systems often suffer from poor retrieval precision due to suboptimal chunking strategies. "
        "Splitting text by fixed character limits often truncates sentences or separates headers from their corresponding paragraphs. "
        "Recent studies advocate for semantic or section-based chunking, which preserves the logical structure of documents. "
        "Furthermore, two-stage retrieval—utilizing a bi-encoder for initial fast retrieval followed by a cross-encoder for precise re-ranking—has become the industry standard for achieving high Mean Reciprocal Rank (MRR). "
        "InterTalks builds upon these findings, implementing section-based chunking and a sophisticated two-stage retrieval pipeline.", justify=True)
    doc.add_page_break()

    # System Architecture
    add_heading(doc, 'Chapter 3: System Architecture', 1)
    add_paragraph(doc, 
        "The InterTalks system employs a modern, decoupled three-tier architecture to ensure scalability, maintainability, and robust performance. "
        "The architecture is logically divided into the Presentation Layer, the Application Layer, and the Data/AI Layer.", justify=True)
    
    # Insert Architecture Diagram
    arch_image = r"C:\Users\nikhil chauhan\.gemini\antigravity\brain\b7eb457a-585b-4a09-b8a8-69f65f1d8467\system_architecture_diagram_1778414290124.png"
    if os.path.exists(arch_image):
        doc.add_picture(arch_image, width=Inches(6.0))
        add_paragraph(doc, "Figure 3.1: InterTalks System Architecture Diagram", align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, '3.1 Presentation Layer (Frontend)', 2)
    add_paragraph(doc, 
        "The frontend is a Single Page Application (SPA) built using React.js and Vite. It utilizes React Router for client-side navigation "
        "and manages global state using Context API, particularly for authentication (AuthContext). "
        "The UI strictly adheres to a dark-themed, enterprise-grade aesthetic, employing custom CSS for layouts, glassmorphism effects, and responsive design. "
        "It communicates with the backend via RESTful APIs using standard HTTP methods.", justify=True)

    add_heading(doc, '3.2 Application Layer (Backend)', 2)
    add_paragraph(doc, 
        "The backend is developed with FastAPI, chosen for its high performance, asynchronous capabilities, and automatic OpenAPI documentation. "
        "The core services include: "
        "\n- AuthService: Handles JWT-based authentication, user registration, and role validation. "
        "\n- DocumentService: Manages document ingestion, file parsing, duplicate detection via checksums, and the approval workflow. "
        "\n- QueryService: Orchestrates the RAG pipeline, including query expansion, multi-stage retrieval, re-ranking, and response generation. "
        "\n- EmbeddingService: Interfaces with local machine learning models to generate dense vectors for documents and queries.", justify=True)

    add_heading(doc, '3.3 Data and AI Layer', 2)
    add_paragraph(doc, 
        "The data tier is bifurcated into structured relational data and unstructured vector data: "
        "\n- PostgreSQL: Acts as the primary source of truth, storing user profiles, document metadata, version histories, and comprehensive audit logs. "
        "\n- Qdrant Vector Database: Stores the 384-dimensional dense vectors generated from document chunks. It is optimized for cosine similarity search and supports payload filtering (e.g., retrieving only 'APPROVED' chunks). "
        "\n- AI Models: The system utilizes 'all-MiniLM-L6-v2' via Sentence Transformers for generating embeddings, and 'ms-marco-MiniLM-L-6-v2' as a Cross-Encoder for precise re-ranking.", justify=True)
    doc.add_page_break()

    # Technology Stack
    add_heading(doc, 'Chapter 4: Technology Stack', 1)
    
    add_heading(doc, '4.1 Frontend Technologies', 2)
    add_paragraph(doc, 
        "- React.js: A declarative, component-based JavaScript library for building user interfaces. "
        "\n- Vite: A next-generation frontend tooling build tool that significantly improves the frontend development experience. "
        "\n- Vanilla CSS: Utilized for styling the application, ensuring maximum flexibility and achieving the desired premium dark-mode aesthetics.", justify=True)

    add_heading(doc, '4.2 Backend Technologies', 2)
    add_paragraph(doc, 
        "- FastAPI: A modern, fast web framework for building APIs with Python based on standard Python type hints. "
        "\n- Uvicorn: An ASGI web server implementation for Python. "
        "\n- SQLAlchemy: The Python SQL toolkit and Object Relational Mapper that provides a full suite of well known enterprise-level persistence patterns. "
        "\n- Pydantic: Used for data parsing and validation, ensuring that data structures are strictly typed.", justify=True)

    add_heading(doc, '4.3 AI & Machine Learning', 2)
    add_paragraph(doc, 
        "- Sentence Transformers: A Python framework for state-of-the-art sentence, text, and image embeddings. "
        "\n- pdfplumber & python-docx: Specialized libraries used to extract text and structure from complex document formats. "
        "\n- Qdrant: A vector similarity search engine and vector database. It provides a production-ready service with a convenient API to store, search, and manage points (vectors) with an additional payload.", justify=True)
    doc.add_page_break()

    # Methodology & Implementation
    add_heading(doc, 'Chapter 5: Methodology & Implementation', 1)
    
    add_heading(doc, '5.1 Document Ingestion Flow', 2)
    add_paragraph(doc, 
        "The ingestion pipeline is designed to be rigorous. When a 'Contributor' uploads a file, it undergoes validation, text extraction, and checksum generation to prevent duplicates. "
        "The document is stored in an 'UPLOADED' state. It requires a 'Reviewer' or 'Admin' to review and transition it to the 'APPROVED' state. "
        "Only upon approval does the system trigger the chunking and embedding processes.", justify=True)
    
    ingest_image = r"C:\Users\nikhil chauhan\.gemini\antigravity\brain\b7eb457a-585b-4a09-b8a8-69f65f1d8467\document_ingestion_flowchart_1778414336142.png"
    if os.path.exists(ingest_image):
        doc.add_picture(ingest_image, width=Inches(6.0))
        add_paragraph(doc, "Figure 5.1: Document Ingestion and Approval Pipeline", align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, '5.2 Section-Based Chunking', 2)
    add_paragraph(doc, 
        "A critical engineering decision in InterTalks is the rejection of arbitrary token-based chunking in favor of semantic section-based chunking. "
        "The system uses regular expressions to detect document headers (e.g., Markdown headers, numbered sections, ALL CAPS titles). "
        "Each chunk represents exactly one logical policy section, preventing the mixing of unrelated topics (cross-policy leakage). "
        "If a section is excessively long (e.g., >600 words), it is sub-split by paragraph while retaining the parent section title.", justify=True)

    add_heading(doc, '5.3 Embedding Strategy', 2)
    add_paragraph(doc, 
        "Before generating the embedding for a chunk, the system prepends the section title to the content (e.g., 'Sick Leave Policy: Employees are entitled to...'). "
        "This provides a much stronger semantic signal for the embedding model, especially for short, targeted queries. "
        "The model 'all-MiniLM-L6-v2' maps these texts to a 384-dimensional vector space. "
        "The vectors, along with extensive metadata (document ID, version, section title, extracted keywords), are upserted into Qdrant.", justify=True)

    add_heading(doc, '5.4 Retrieval-Augmented Generation Pipeline', 2)
    add_paragraph(doc, 
        "The query pipeline features several stages to maximize accuracy: "
        "\n1. Query Preprocessing & Expansion: Short queries (<4 words) are automatically expanded into multiple variants to improve matching. "
        "\n2. Initial Retrieval: A vector search is performed in Qdrant, returning the top 10 candidates. A strict filter ensures only 'APPROVED' documents are searched. "
        "\n3. Composite Scoring: Candidates are re-ranked using a weighted formula: 40% Semantic Similarity + 30% Keyword Match + 20% Title Match + 10% Exact Match. "
        "\n4. Adaptive Thresholding: The system applies a dynamic minimum similarity threshold based on query length (e.g., 0.40 for short, 0.60 for long). "
        "\n5. Cross-Encoder Re-Ranking: The top candidates are processed through a highly accurate cross-encoder model to determine the final, best context.", justify=True)

    rag_image = r"C:\Users\nikhil chauhan\.gemini\antigravity\brain\b7eb457a-585b-4a09-b8a8-69f65f1d8467\rag_pipeline_flowchart_1778414304850.png"
    if os.path.exists(rag_image):
        doc.add_picture(rag_image, width=Inches(6.0))
        add_paragraph(doc, "Figure 5.2: RAG Pipeline Flowchart", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # Database Schema
    add_heading(doc, 'Chapter 6: Database Schema', 1)
    add_paragraph(doc, 
        "The PostgreSQL database employs a normalized schema to manage the complex relationships between users, documents, versions, and audits. "
        "Key tables include: "
        "\n- users: Stores authentication details, roles, and status. "
        "\n- documents: Represents the logical container for a policy. "
        "\n- document_versions: Stores distinct versions of a document, tracking their lifecycle state (UPLOADED, UNDER_REVIEW, APPROVED, DEPRECATED) and checksums. "
        "\n- chunks: Contains the parsed text segments, section numbers, and token counts. "
        "\n- audit_logs & query_audit_logs: Maintain an immutable trail of system actions, crucial for compliance and monitoring.", justify=True)
    doc.add_page_break()

    # User Interface
    add_heading(doc, 'Chapter 7: User Interface & User Experience', 1)
    add_paragraph(doc, 
        "The User Interface is crafted to provide a seamless, intuitive experience for different roles. "
        "It features a consistent dark theme, emphasizing content readability and reduced eye strain. "
        "Key interfaces include the Login/Registration portals, role-specific Dashboards, the Document Upload interface, and the Ask Query page.", justify=True)
    
    # Insert Login UI
    login_image = r"C:\Users\nikhil chauhan\.gemini\antigravity\brain\b7eb457a-585b-4a09-b8a8-69f65f1d8467\login_page_ui_1778414198644.png"
    if os.path.exists(login_image):
        doc.add_picture(login_image, width=Inches(5.5))
        add_paragraph(doc, "Figure 7.1: InterTalks Login Portal", align=WD_ALIGN_PARAGRAPH.CENTER)

    # Insert Admin UI
    admin_image = r"C:\Users\nikhil chauhan\.gemini\antigravity\brain\b7eb457a-585b-4a09-b8a8-69f65f1d8467\admin_dashboard_ui_1778414212122.png"
    if os.path.exists(admin_image):
        doc.add_picture(admin_image, width=Inches(5.5))
        add_paragraph(doc, "Figure 7.2: Administrator Dashboard", align=WD_ALIGN_PARAGRAPH.CENTER)

    # Insert Upload UI
    upload_image = r"C:\Users\nikhil chauhan\.gemini\antigravity\brain\b7eb457a-585b-4a09-b8a8-69f65f1d8467\upload_page_ui_1778414254476.png"
    if os.path.exists(upload_image):
        doc.add_picture(upload_image, width=Inches(5.5))
        add_paragraph(doc, "Figure 7.3: Document Upload Interface", align=WD_ALIGN_PARAGRAPH.CENTER)

    # Insert Ask UI
    ask_image = r"C:\Users\nikhil chauhan\.gemini\antigravity\brain\b7eb457a-585b-4a09-b8a8-69f65f1d8467\ask_query_page_ui_1778414239064.png"
    if os.path.exists(ask_image):
        doc.add_picture(ask_image, width=Inches(5.5))
        add_paragraph(doc, "Figure 7.4: Ask Query Page with Grounded Answer", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # Conclusion
    add_heading(doc, 'Chapter 8: Conclusion & Future Enhancements', 1)
    add_paragraph(doc, 
        "The InterTalks project successfully demonstrates a highly controlled, deterministic RAG system tailored for corporate compliance. "
        "By enforcing strict ingestion rules, utilizing section-based chunking, and implementing a multi-stage retrieval pipeline with cross-encoder re-ranking, "
        "the system significantly mitigates the risk of AI hallucinations. The robust governance model ensures that users only interact with approved, verified data. "
        "\n\nFuture enhancements may include: "
        "\n- Multi-document synthesis for queries spanning several policies. "
        "\n- Advanced OCR capabilities for scanned, non-text-searchable PDFs. "
        "\n- Active learning mechanisms to tune the composite scoring weights based on user feedback.", justify=True)

    doc.save('InterTalks_Project_Report.docx')

if __name__ == '__main__':
    create_report()
