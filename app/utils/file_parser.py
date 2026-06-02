import hashlib
import io
import re
from typing import Tuple, List, Dict, Any, Set
from fastapi import UploadFile

# Try to import Document parser dependencies
try:
    from docx import Document
except ImportError:
    Document = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

def generate_checksum(text: str) -> str:
    """Generates a SHA-256 checksum for the exact text."""
    return hashlib.sha256(text.encode('utf-8')).hexdigest()

def normalize_text(text: str) -> str:
    """Removes extra whitespaces, normalize returns."""
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'[\n\r]{2,}', '\n\n', text)
    return text.strip()


# ── Text Cleaning ──────────────────────────────────────────────────

def clean_text(text: str) -> str:
    """
    Production text cleaner — removes noise BEFORE chunking.
    Strips: page breaks, page numbers, repeated header/footer lines,
    excessive whitespace, and common PDF artifacts.
    """
    # Remove PAGE BREAK markers
    text = re.sub(r'---\s*PAGE BREAK\s*---', '', text)

    # Remove standalone page numbers (e.g. "Page 3", "- 3 -", "3")
    text = re.sub(r'(?m)^\s*(?:Page\s+)?\d{1,4}\s*$', '', text, flags=re.IGNORECASE)
    text = re.sub(r'(?m)^\s*-\s*\d{1,4}\s*-\s*$', '', text)

    # Remove common PDF header/footer patterns
    text = re.sub(r'(?m)^.*(?:CONFIDENTIAL|INTERNAL USE ONLY|DRAFT).*$', '', text, flags=re.IGNORECASE)

    # Remove repeated whitespace-only lines
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Strip leading/trailing whitespace per line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)

    # Remove duplicate consecutive lines (header/footer repetition)
    seen_lines = []
    deduped = []
    for line in text.split('\n'):
        stripped = line.strip()
        if stripped and stripped in seen_lines[-3:] if seen_lines else False:
            continue  # Skip repeated line
        seen_lines.append(stripped)
        deduped.append(line)
    text = '\n'.join(deduped)

    return text.strip()


# ── Keyword Extraction ─────────────────────────────────────────────

_STOP_WORDS: Set[str] = {
    'a', 'an', 'the', 'is', 'are', 'was', 'were', 'be', 'been', 'being',
    'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could',
    'should', 'may', 'might', 'shall', 'can', 'must', 'need',
    'and', 'or', 'but', 'if', 'then', 'else', 'when', 'where', 'how',
    'what', 'which', 'who', 'whom', 'this', 'that', 'these', 'those',
    'to', 'of', 'in', 'on', 'at', 'by', 'for', 'with', 'from', 'about',
    'into', 'through', 'during', 'before', 'after', 'above', 'below',
    'up', 'down', 'out', 'off', 'over', 'under', 'again', 'further',
    'not', 'no', 'nor', 'so', 'too', 'very', 'just', 'also', 'only',
    'such', 'than', 'both', 'each', 'all', 'any', 'few', 'more', 'most',
    'other', 'some', 'its', 'own', 'same', 'as', 'per', 'etc',
    'i', 'me', 'my', 'we', 'our', 'you', 'your', 'he', 'she', 'it',
    'they', 'them', 'their', 'his', 'her'
}


def extract_keywords(section_title: str, content: str) -> List[str]:
    """
    Extract meaningful keywords from section title + content.
    Returns deduplicated list of keywords (lowercased).
    """
    # Title keywords are HIGH priority — always include
    title_words = set(re.findall(r'[a-zA-Z]{2,}', section_title.lower())) - _STOP_WORDS

    # Content keywords — extract significant words
    content_words_raw = re.findall(r'[a-zA-Z]{3,}', content.lower())
    content_words = [w for w in content_words_raw if w not in _STOP_WORDS]

    # Count frequency for content words
    word_freq = {}
    for w in content_words:
        word_freq[w] = word_freq.get(w, 0) + 1

    # Take top 15 most frequent content words
    top_content = sorted(word_freq.keys(), key=lambda w: word_freq[w], reverse=True)[:15]

    # Combine: title words first, then top content words
    all_keywords = list(title_words)
    for w in top_content:
        if w not in title_words:
            all_keywords.append(w)

    # Also add bigrams from title (e.g. "sick leave" → "sick leave")
    title_clean = re.sub(r'[^a-zA-Z\s]', '', section_title.lower()).strip()
    if title_clean and len(title_clean.split()) >= 2:
        all_keywords.insert(0, title_clean)

    return all_keywords[:20]  # Cap at 20 keywords


# ── Document Extraction ────────────────────────────────────────────

def extract_docx(file_bytes: bytes) -> str:
    """Extracts text from DOCX preserving paragraph structure."""
    if not Document:
        raise ImportError("python-docx is not installed")

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            if para.style.name.startswith('Heading'):
                level = para.style.name.replace('Heading ', '')
                try:
                    level_int = max(1, min(6, int(level)))
                    paragraphs.append(f"{'#' * level_int} {text}")
                except ValueError:
                    paragraphs.append(f"# {text}")
            else:
                paragraphs.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_items = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_items:
                paragraphs.append(" | ".join(row_items))

    raw_text = "\n\n".join(paragraphs)
    return normalize_text(raw_text)


def extract_pdf(file_path: str) -> str:
    """Extracts text from PDF preserving logical reading order."""
    if not pdfplumber:
        raise ImportError("pdfplumber is not installed")

    text_content = []

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text(layout=False)
            if page_text:
                lines = page_text.split('\n')
                if len(lines) > 10:
                    lines = lines[1:-1]
                text_content.append('\n'.join(lines))

    raw_text = "\n\n".join(text_content)
    return normalize_text(raw_text)


def estimate_tokens(text: str) -> int:
    """Rough token estimator (words × 1.3)."""
    words = len(re.findall(r'\w+', text))
    return int(words * 1.3)


# ── Section Heading Detection ──────────────────────────────────────

_MD_HEADING_RE = re.compile(r'^(#{1,6})\s+(.+)$')
_NUMBERED_HEADING_RE = re.compile(r'^(\d+(?:\.\d+)*)\s+(.+)$')
_ALLCAPS_HEADING_RE = re.compile(r'^([A-Z][A-Z\s]{8,})$')


def _detect_heading(line: str):
    """
    Detect if a line is a section heading.
    Returns (section_number, section_title) or None.
    """
    line = line.strip()
    if not line:
        return None

    match = _MD_HEADING_RE.match(line)
    if match:
        return (None, match.group(2).strip())

    match = _NUMBERED_HEADING_RE.match(line)
    if match:
        return (match.group(1), match.group(2).strip())

    match = _ALLCAPS_HEADING_RE.match(line)
    if match:
        return (None, line.strip().title())

    return None


# ── Section-Based Chunking (Production) ────────────────────────────

def split_into_sections(text: str, max_section_words: int = 600) -> List[Dict[str, Any]]:
    """
    Production section-based chunking.

    Rules:
      - CLEAN text first (remove noise)
      - Each chunk = exactly ONE logical policy section
      - ONE chunk = ONE intent (never merge topics)
      - Each chunk includes: section_number, section_title, content, keywords, token_count
      - Oversized sections sub-split by paragraph
    """
    # ── Step 0: Clean text ──────────────────────────────────────
    text = clean_text(text)

    lines = text.split('\n')

    # ── Pass 1: Find section boundaries ─────────────────────────
    sections = []
    current_title = None
    current_number = None
    current_lines = []
    auto_section_counter = 0

    for line in lines:
        heading = _detect_heading(line)

        if heading is not None:
            # Flush previous section
            if current_lines:
                body = '\n'.join(current_lines).strip()
                if body:
                    auto_section_counter += 1
                    sections.append({
                        'section_number': current_number or str(auto_section_counter),
                        'section_title': current_title or f'Section {auto_section_counter}',
                        'content': body
                    })

            current_number, current_title = heading
            current_lines = []
        else:
            current_lines.append(line)

    # Flush last section
    if current_lines:
        body = '\n'.join(current_lines).strip()
        if body:
            auto_section_counter += 1
            sections.append({
                'section_number': current_number or str(auto_section_counter),
                'section_title': current_title or f'Section {auto_section_counter}',
                'content': body
            })

    # ── Filter: skip preamble/metadata-only sections ────────────
    # Preamble chunks (e.g. "Version 1.0, Department: HR") contain no
    # actual policy content and pollute search results.
    # Skip sections that are very short AND have a generic auto-title.
    filtered_sections = []
    for section in sections:
        word_count = len(section['content'].split())
        is_generic_title = section['section_title'].startswith('Section ')
        # Skip if < 30 words AND has a generic auto-generated title
        if word_count < 30 and is_generic_title:
            continue
        filtered_sections.append(section)

    # If filtering removed everything, keep the originals
    if filtered_sections:
        sections = filtered_sections

    # ── Edge case: no headings → split by paragraphs ────────────
    if not sections and text.strip():
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        if len(paragraphs) <= 1:
            sections = [{
                'section_number': '1',
                'section_title': 'Document Content',
                'content': text.strip()
            }]
        else:
            for i, para in enumerate(paragraphs, 1):
                first_line = para.split('\n')[0].strip()
                sections.append({
                    'section_number': str(i),
                    'section_title': first_line[:80] if len(first_line) < 100 else f'Section {i}',
                    'content': para
                })

    # ── Pass 2: Sub-split oversized sections + add keywords ─────
    final_chunks = []

    for section in sections:
        word_count = len(section['content'].split())
        title = section['section_title']
        number = section['section_number']

        if word_count <= max_section_words:
            keywords = extract_keywords(title, section['content'])
            final_chunks.append({
                'section_number': number,
                'section_title': title,
                'content': section['content'],
                'keywords': keywords,
                'token_count': estimate_tokens(section['content'])
            })
        else:
            # Sub-split by paragraphs
            paragraphs = [p.strip() for p in section['content'].split('\n\n') if p.strip()]
            sub_chunk_text = ""
            sub_index = 1

            for para in paragraphs:
                para_words = len(para.split())
                current_words = len(sub_chunk_text.split()) if sub_chunk_text else 0

                if current_words + para_words > max_section_words and sub_chunk_text:
                    keywords = extract_keywords(title, sub_chunk_text)
                    final_chunks.append({
                        'section_number': f"{number}.{sub_index}",
                        'section_title': title,
                        'content': sub_chunk_text.strip(),
                        'keywords': keywords,
                        'token_count': estimate_tokens(sub_chunk_text.strip())
                    })
                    sub_index += 1
                    sub_chunk_text = para
                else:
                    sub_chunk_text += ("\n\n" if sub_chunk_text else "") + para

            if sub_chunk_text.strip():
                keywords = extract_keywords(title, sub_chunk_text)
                final_chunks.append({
                    'section_number': f"{number}.{sub_index}" if sub_index > 1 else number,
                    'section_title': title,
                    'content': sub_chunk_text.strip(),
                    'keywords': keywords,
                    'token_count': estimate_tokens(sub_chunk_text.strip())
                })

    return final_chunks


# Backward-compatible alias
def split_into_chunks(text: str, max_chunk_words: int = 400, overlap_words: int = 50) -> List[Dict[str, Any]]:
    """Backward-compatible wrapper — delegates to section-based splitting."""
    return split_into_sections(text)
